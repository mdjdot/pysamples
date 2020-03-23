#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches
from docx.shared import Cm

"""
https://python-docx.readthedocs.io/en/latest/index.html
"""


def writeWord():
    doc = Document()

    doc.add_heading("1. the first heading", level=1)
    doc.add_heading("1.1 the sub heading", level=2)

    doc.add_paragraph("some sentences")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Age"
    row = table.rows[1]
    row.cells[0].text = "张三"
    row.cells[1].text = "18"
    table.style = "LightShading-Accent1"

    doc.add_picture("./timg.jpg", width=Inches(2))

    doc.add_page_break()

    doc.add_heading("2. the second heading", level=1)
    doc.add_heading("2.1 the sub heading", level=2)

    doc.add_paragraph("some sentences with style", style="ListBullet")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Age"
    row = table.rows[1]
    row.cells[0].text = "张三"
    row.cells[1].text = "18"
    table.style = "LightShading-Accent1"

    doc.add_picture("./timg.jpg", width=Cm(10))

    doc.save("sample.docx")


def readWord():
    doc = Document("./sample.docx")
    for para in doc.paragraphs:
        print(para.text)

    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                print(cell.text)


if __name__ == "__main__":
    writeWord()
    readWord()
